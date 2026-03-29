import fitz  # PyMuPDF — import name is "fitz", not "pymupdf"
import docx
from pathlib import Path
from typing import List, Dict, Any

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma

from app.utils import (
    CHROMA_DB_PATH, CHUNK_SIZE, CHUNK_OVERLAP, OPENAI_API_KEY
)

def load_pdf(file_path: str) -> List[Document]:
    """
    Extract text from a PDF using PyMuPDF.
    Returns a list of LangChain Document objects, one per page.
    Each Document has page_content (the text) and metadata (source, page).
    """
    documents = []
    # fitz.open() loads the PDF into memory
    pdf = fitz.open(file_path)

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        # get_text("text") extracts plain text, preserving layout
        text = page.get_text("text")

        # Skip pages with no text (scanned images, blank pages)
        if not text.strip():
            continue

        documents.append(Document(
            page_content=text,
            metadata={
                "source": Path(file_path).name,
                "page": page_num + 1,  # 1-indexed for humans
                "file_path": str(file_path),
                "file_type": "pdf",
            }
        ))

    pdf.close()
    return documents

def load_docx(file_path: str) -> List[Document]:
    """
    Extract text from a Word document.
    DOCX files don't have "pages" in a predictable way, so we treat
    the whole document as one Document and chunk it later.
    """
    doc = docx.Document(file_path)

    # Each paragraph is a separate line in the Word doc
    full_text = "\n".join([para.text for para in doc.paragraphs])

    return [Document(
        page_content=full_text,
        metadata={
            "source": Path(file_path).name,
            "page": 1,
            "file_path": str(file_path),
            "file_type": "docx",
        }
    )]

def load_txt(file_path: str) -> List[Document]:
    """Load a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    return [Document(
        page_content=text,
        metadata={
            "source": Path(file_path).name,
            "page": 1,
            "file_path": str(file_path),
            "file_type": "txt",
        }
    )]

def load_document(file_path: str) -> List[Document]:
    """
    Router function: detect file type and call the right loader.
    Raises ValueError for unsupported file types.
    """
    ext = Path(file_path).suffix.lower()

    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".doc": load_docx,
        ".txt": load_txt,
        ".md": load_txt,
    }

    if ext not in loaders:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {list(loaders.keys())}"
        )

    return loaders[ext](file_path)

def chunk_documents(
    documents: List[Document],
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> List[Document]:
    """
    Split documents into chunks using RecursiveCharacterTextSplitter.

    How RecursiveCharacterTextSplitter works:
    It tries to split on each separator in order, stopping when chunks
    are small enough. So it tries "\n\n" first (paragraph breaks),
    then "\n" (line breaks), then ". " (sentence ends), then " " (words),
    then "" (characters). This means chunks respect natural language
    boundaries as much as possible.

    chunk_size: max characters per chunk (500 chars ≈ ~125 tokens)
    chunk_overlap: how many characters to repeat between adjacent chunks
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
        # length_function counts characters, not tokens.
        # For token-accurate chunking, swap with:
        # length_function=lambda t: len(tiktoken.get_encoding("cl100k_base").encode(t))
        length_function=len,
    )

    chunks = splitter.split_documents(documents)

    # Add chunk index to metadata for debugging
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
        chunk.metadata["chunk_total"] = len(chunks)

    return chunks

def get_vectorstore() -> Chroma:
    """
    Get or create the ChromaDB vector store.
    This is a singleton-like function — calling it multiple times
    returns the same persisted collection on disk.
    """
    embeddings = OpenAIEmbeddings(
        model="text-embedding-3-small",
        openai_api_key=OPENAI_API_KEY,
    )
    return Chroma(
        collection_name="research_docs",
        embedding_function=embeddings,
        persist_directory=CHROMA_DB_PATH,
    )

def ingest_document(file_path: str) -> Dict[str, Any]:
    """
    Full ingestion pipeline for one document.
    Returns a summary dict with ingestion stats.

    Steps:
    1. Load the document (extract text)
    2. Chunk the text
    3. Embed each chunk + store in ChromaDB
    4. Return stats
    """
    print(f"[Ingestion] Loading: {file_path}")
    documents = load_document(file_path)
    print(f"[Ingestion] Loaded {len(documents)} pages")

    chunks = chunk_documents(documents)
    print(f"[Ingestion] Created {len(chunks)} chunks")

    # Check if this document was already ingested (avoid duplicates)
    filename = Path(file_path).name
    vectorstore = get_vectorstore()
    existing = vectorstore.get(where={"source": filename})
    if existing["ids"]:
        print(f"[Ingestion] Document already exists, deleting old version")
        # Delete old chunks for this file before re-ingesting
        vectorstore.delete(ids=existing["ids"])

    # This call embeds all chunks and stores them.
    # It makes one API call per chunk to OpenAI's embedding endpoint.
    # With 100 chunks, that's 100 API calls (batched automatically).
    vectorstore.add_documents(chunks)
    print(f"[Ingestion] Stored {len(chunks)} chunks in ChromaDB")

    return {
        "filename": filename,
        "pages": len(documents),
        "chunks": len(chunks),
        "status": "success",
    }

def list_ingested_documents() -> List[str]:
    """Return unique filenames of all ingested documents."""
    vectorstore = get_vectorstore()
    all_metadata = vectorstore.get()["metadatas"]
    filenames = list({m["source"] for m in all_metadata if m})
    return sorted(filenames)